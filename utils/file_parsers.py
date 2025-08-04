import os
import tempfile
from typing import Optional
import PyPDF2
import pdfplumber
from docx import Document
from config import ALLOWED_EXTENSIONS

def parse_file(file_path: str) -> str:
    """
    Parse a file and extract text content.
    
    Args:
        file_path (str): Path to the file
        
    Returns:
        str: Extracted text content
        
    Raises:
        ValueError: If file type is not supported
    """
    # Get file extension
    _, ext = os.path.splitext(file_path)
    ext = ext.lower()[1:]  # Remove the dot
    
    if ext not in ALLOWED_EXTENSIONS:
        raise ValueError(f"Unsupported file type: {ext}")
    
    if ext == "pdf":
        return parse_pdf(file_path)
    elif ext == "docx":
        return parse_docx(file_path)
    elif ext == "txt":
        return parse_txt(file_path)
    else:
        raise ValueError(f"Unsupported file type: {ext}")

def parse_pdf(file_path: str) -> str:
    """
    Parse a PDF file and extract text content.
    
    Args:
        file_path (str): Path to the PDF file
        
    Returns:
        str: Extracted text content
    """
    text = ""
    
    try:
        # First try with PyPDF2
        with open(file_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            for page in reader.pages:
                text += page.extract_text() + "\n"
        
        # If PyPDF2 didn't extract much text, try with pdfplumber
        if len(text.strip()) < 100:
            text = ""
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
    except Exception as e:
        raise ValueError(f"Error parsing PDF file: {str(e)}")
    
    return text

def parse_docx(file_path: str) -> str:
    """
    Parse a DOCX file and extract text content.
    
    Args:
        file_path (str): Path to the DOCX file
        
    Returns:
        str: Extracted text content
    """
    try:
        doc = Document(file_path)
        text = ""
        
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + "\t"
                text += "\n"
        
        return text
    except Exception as e:
        raise ValueError(f"Error parsing DOCX file: {str(e)}")

def parse_txt(file_path: str) -> str:
    """
    Parse a TXT file and extract text content.
    
    Args:
        file_path (str): Path to the TXT file
        
    Returns:
        str: Extracted text content
    """
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()
    except UnicodeDecodeError:
        # Try with different encoding
        try:
            with open(file_path, 'r', encoding='latin-1') as file:
                return file.read()
        except Exception as e:
            raise ValueError(f"Error parsing TXT file: {str(e)}")
    except Exception as e:
        raise ValueError(f"Error parsing TXT file: {str(e)}")